"""
redact_docx.py — redacts the actual .docx file in place, editing text at
the RUN level (not full paragraph rewrite) so bold/fonts/tables/layout
survive. Word often splits one visible phrase across multiple runs, so we
join all run text in a paragraph/cell first (full context for detection),
then map matches back onto the original runs.
"""

import json
from pathlib import Path
from typing import List

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.detector import detect
from src.fake_mapper import FakeMapper
from src.models import Match

mapper = FakeMapper()


def _redact_runs(runs: List, log: List[dict]) -> None:
    """Generic run-list redactor. Works for a single paragraph's runs
    or for all runs flattened across every paragraph in a table cell."""
    if not runs:
        return

    full_text = " ".join(r.text for r in runs)
    if not full_text.strip():
        return

    matches: List[Match] = detect(full_text)
    if not matches:
        return

    run_spans = []
    cursor = 0
    for r in runs:
        run_spans.append((cursor, cursor + len(r.text)))
        cursor += len(r.text) + 1  # +1 accounts for the " " separator

    for m in sorted(matches, key=lambda m: m.start, reverse=True):
        fake = mapper.get(m.text, m.label)
        log.append({"type": m.label, "original": m.text, "replacement": fake,
                     "confidence": m.confidence, "source": m.source, "rule": m.rule_name})

        touched = [i for i, (s, e) in enumerate(run_spans) if not (m.end <= s or m.start >= e)]
        if not touched:
            continue

        first_idx, last_idx = touched[0], touched[-1]
        first_start, _ = run_spans[first_idx]
        last_start, _ = run_spans[last_idx]
        prefix = runs[first_idx].text[: m.start - first_start]

        if first_idx == last_idx:
            suffix = runs[first_idx].text[m.end - first_start:]
            runs[first_idx].text = prefix + fake + suffix
        else:
            runs[first_idx].text = prefix + fake
            for i in touched[1:-1]:
                runs[i].text = ""
            suffix = runs[last_idx].text[m.end - last_start:]
            runs[last_idx].text = suffix


def _redact_paragraph_runs(paragraph: Paragraph, log: List[dict]) -> None:
    """Used for regular body paragraphs (outside tables), headers, footers."""
    _redact_runs(paragraph.runs, log)


def _redact_table(table: Table, log: List[dict]) -> None:
    for row in table.rows:
        seen_cell_ids = set()
        for cell in row.cells:
            if id(cell) in seen_cell_ids:
                continue  # merged cell already processed via an earlier column
            seen_cell_ids.add(id(cell))
            all_runs = [r for para in cell.paragraphs for r in para.runs]
            _redact_runs(all_runs, log)
            for nested in cell.tables:
                _redact_table(nested, log)


def redact_docx(input_path: str, output_path: str) -> List[dict]:
    doc = Document(input_path)
    log: List[dict] = []

    for para in doc.paragraphs:
        _redact_paragraph_runs(para, log)
    for table in doc.tables:
        _redact_table(table, log)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _redact_paragraph_runs(para, log)
        for para in section.footer.paragraphs:
            _redact_paragraph_runs(para, log)
        for table in section.header.tables:
            _redact_table(table, log)
        for table in section.footer.tables:
            _redact_table(table, log)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return log